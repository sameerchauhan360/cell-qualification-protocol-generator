import re

import docx


def parse_tmp_document(tmp_path):
    doc = docx.Document(tmp_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if len(paragraphs) < 2:
        raise ValueError(f"TMP document {tmp_path} does not contain enough paragraphs.")

    tmp_doc_title = paragraphs[0]
    second_para = paragraphs[1]
    parts = [part.strip() for part in second_para.split("|")]

    chemistry = ""
    if len(parts) >= 3:
        chem_part = parts[2]
        chem_part = re.split(r"(?i)prepared\s+under", chem_part)[0].strip()
        chemistry = chem_part

    return {"tmp_doc_title": tmp_doc_title, "chemistry": chemistry}


def parse_acl_document(acl_path):
    doc = docx.Document(acl_path)
    paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if len(paragraphs_text) < 2:
        raise ValueError(f"ACL document {acl_path} does not contain enough paragraphs.")

    acl_doc_title = paragraphs_text[0]
    second_para = paragraphs_text[1]
    parts = [part.strip() for part in second_para.split("|")]
    doc_number = ""
    for part in parts:
        if "Document" in part:
            doc_code = part.replace("Document", "").strip()
            doc_number = doc_code.replace("ACL-", "CQP-")
            break

    duty_profiles = []
    footnotes = []
    elements = doc.element.body
    current_profile = None
    tables_map = {tbl._tbl: tbl for tbl in doc.tables}
    notes_started = False

    for child in elements:
        tag = child.tag
        if tag.endswith("p"):
            text = "".join(
                t.text
                for t in child.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                )
                if t.text
            ).strip()
            if not text:
                continue

            if text.startswith("Duty Profile:"):
                profile_match = re.match(
                    r"Duty Profile:\s*([^\(]+)(?:\(conditioning rates:\s*([^\)]+)\))?",
                    text,
                )
                if profile_match:
                    profile_name = profile_match.group(1).strip()
                    rates_str = profile_match.group(2)
                    rates = []
                    if rates_str:
                        rates = [r.strip() for r in rates_str.split(",")]

                    current_profile = {
                        "name": profile_name,
                        "rates": rates,
                        "tests": [],
                        "cycles": 500,
                    }
                    duty_profiles.append(current_profile)

            elif text.startswith("Notes:") or text.startswith("Notes"):
                notes_started = True
            elif notes_started:
                if any(text.startswith(marker) for marker in ["*", "#", "@", "$"]):
                    footnotes.append(text)

        elif tag.endswith("tbl"):
            if current_profile is not None:
                table_obj = tables_map.get(child)
                if table_obj:
                    for row in table_obj.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) >= 4 and cells[1]:
                            sr_no = cells[0]
                            test_param = cells[1]
                            limit = cells[2]
                            clause = cells[3]

                            current_profile["tests"].append(
                                {
                                    "sr_no": sr_no,
                                    "parameter": test_param,
                                    "limit": limit,
                                    "clause": clause,
                                }
                            )

                            if (
                                "cycle-life" in test_param.lower()
                                or "cycle life" in test_param.lower()
                            ):
                                cycles_match = re.search(
                                    r"(\d+)\s*cycles", limit, re.IGNORECASE
                                )
                                if cycles_match:
                                    current_profile["cycles"] = int(
                                        cycles_match.group(1)
                                    )

                current_profile = None

    return {
        "acl_doc_title": acl_doc_title,
        "doc_number": doc_number,
        "duty_profiles": duty_profiles,
        "footnotes": footnotes,
    }


if __name__ == "__main__":
    import json

    base_dir = r"d:\Agentic Ai\Assignment\AIMl Engineer-Webosphere\candidate_pack\candidate_pack\inputs\setA"
    print("Testing TMP parser...")
    tmp_data = parse_tmp_document(f"{base_dir}\\TMP_CYG-21700-50G.docx")
    print(json.dumps(tmp_data, indent=2))

    print("\nTesting ACL parser...")
    acl_data = parse_acl_document(f"{base_dir}\\ACL_CYG-21700-50G.docx")
    print(json.dumps(acl_data, indent=2))
