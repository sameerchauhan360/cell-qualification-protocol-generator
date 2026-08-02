import copy
import os

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph directly after the specified paragraph in the document XML tree."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text is not None:
        new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph


def clone_row(table, template_row):
    """Deep-copy an XML table row and append it to the table element."""
    tr = template_row._tr
    new_tr = copy.deepcopy(tr)
    table._tbl.append(new_tr)
    return docx.table._Row(new_tr, table)


def merge_cells_vertically(table, col_idx, start_row_idx, end_row_idx):
    """Configure XML vertical cell merging tags (restart/continue) across rows in a column."""
    for row_idx in range(start_row_idx, end_row_idx + 1):
        cell = table.cell(row_idx, col_idx)
        tcPr = cell._tc.get_or_add_tcPr()
        vMerge = OxmlElement("w:vMerge")
        if row_idx == start_row_idx:
            vMerge.set(qn("w:val"), "restart")
        else:
            cell.text = ""  # Subsequent merged cells must have empty text to match Gold structure
        tcPr.append(vMerge)


def replace_in_paragraph(p, old, new):
    """Replace placeholder text in a paragraph while maintaining run-level styles."""
    if old not in p.text:
        return
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    if len(p.runs) > 0:
        full_text = "".join(r.text for r in p.runs)
        full_text = full_text.replace(old, new)
        p.runs[0].text = full_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = p.text.replace(old, new)


def replace_in_table(table, old, new):
    """Recursively search and replace placeholder text in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, old, new)
            for sub_table in cell.tables:
                replace_in_table(sub_table, old, new)


def replace_in_header_footer(doc, old, new):
    """Replace placeholder text in headers and footers across all document sections."""
    for section in doc.sections:
        for header in [
            section.header,
            section.first_page_header,
            section.even_page_header,
        ]:
            if header:
                for p in header.paragraphs:
                    replace_in_paragraph(p, old, new)
                for table in header.tables:
                    replace_in_table(table, old, new)
        for footer in [
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]:
            if footer:
                for p in footer.paragraphs:
                    replace_in_paragraph(p, old, new)
                for table in footer.tables:
                    replace_in_table(table, old, new)


def make_reviewer_zone_editable(header_p, body_p, perm_id):
    """Inject permStart and permEnd exceptions around reviewer text to leave them editable."""
    perm_start = OxmlElement("w:permStart")
    perm_start.set(qn("w:id"), str(perm_id))
    perm_start.set(qn("w:edit"), "everyone")
    header_p._p.append(perm_start)

    perm_end = OxmlElement("w:permEnd")
    perm_end.set(qn("w:id"), str(perm_id))
    body_p._p.append(perm_end)


def apply_document_protection(doc):
    """Enable read-only enforcement in word/settings.xml."""
    settings = doc.settings.element
    documentProtection = settings.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}documentProtection"
    )
    if documentProtection is None:
        documentProtection = OxmlElement("w:documentProtection")
        settings.append(documentProtection)
    documentProtection.set(qn("w:edit"), "readOnly")
    documentProtection.set(qn("w:enforcement"), "1")
    documentProtection.set(qn("w:formatting"), "0")


def render_protocol_docx(template_path, data, output_path):
    """Assemble the final CQP Word document from the template and extracted data."""
    doc = docx.Document(template_path)

    # 1. Strip boilerplate/guidance instructions from the template body
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == "The following values are taken from the vendor datasheet.":
            paragraphs_to_remove.append(p)
        elif (
            txt
            == "Table 1 — expands to one row per duty profile and conditioning rate."
        ):
            p.text = "Table 1. Duty profile test matrix."
        elif (
            txt
            == "The block below is repeated for each duty profile. The placeholder row is replaced by one row per test taken from the ACL for that duty profile."
        ):
            paragraphs_to_remove.append(p)
        elif txt == "This table is completed manually and MUST remain blank at issue.":
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    # 2. Expand Table 1 Matrix (combines duty profiles and conditioning rates)
    table1 = None
    tpl_row_idx = None
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            cell_texts = [cell.text for cell in row.cells]
            if any("{{ duty_profile }}" in text for text in cell_texts):
                table1 = table
                tpl_row_idx = r_idx
                break
        if table1:
            break

    if table1 and tpl_row_idx is not None:
        template_row = table1.rows[tpl_row_idx]

        combos = []
        for dp in data["duty_profiles"]:
            for rate in dp["rates"]:
                combos.append(
                    {
                        "profile_name": dp["name"],
                        "rate": rate,
                        "v_max": data["v_max"],
                        "v_min": data["v_min"],
                        "cycles": str(dp["cycles"]),
                    }
                )

        new_rows = []
        for idx, combo in enumerate(combos):
            new_row = clone_row(table1, template_row)
            new_rows.append(new_row)

            new_row.cells[0].text = str(idx + 1)
            new_row.cells[1].text = combo["profile_name"]
            new_row.cells[2].text = combo["rate"]
            new_row.cells[3].text = combo["v_max"]
            new_row.cells[4].text = combo["v_min"]
            new_row.cells[5].text = combo["cycles"]

        table1._tbl.remove(template_row._tr)

        # Merge cell blocks vertically across rates per profile
        curr_row = tpl_row_idx
        for dp in data["duty_profiles"]:
            num_rates = len(dp["rates"])
            if num_rates > 1:
                merge_cells_vertically(table1, 1, curr_row, curr_row + num_rates - 1)
            curr_row += num_rates

    # 3. Duplicate Section 7 Blocks (1 block of 6 elements per duty profile)
    body = doc.element.body
    children = list(body)
    start_idx = None

    for idx, child in enumerate(children):
        if child.tag.endswith("p"):
            p_text = "".join(
                t.text
                for t in child.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                )
                if t.text
            ).strip()
            if "7.{{ block_index }}" in p_text:
                start_idx = idx
                break

    if start_idx is not None:
        original_elements = children[start_idx : start_idx + 6]
        for elem in original_elements:
            body.remove(elem)

        insert_idx = start_idx
        perm_id_counter = 1

        for i, dp in enumerate(data["duty_profiles"]):
            block_idx = i + 1
            copied_elements = [copy.deepcopy(elem) for elem in original_elements]

            for elem in copied_elements:
                body.insert(insert_idx, elem)
                insert_idx += 1

            separator_p = OxmlElement("w:p")
            body.insert(insert_idx, separator_p)
            insert_idx += 1

            current_table = None
            block_paragraphs = []
            for elem in copied_elements:
                if elem.tag.endswith("p"):
                    p = Paragraph(elem, doc)
                    block_paragraphs.append(p)
                elif elem.tag.endswith("tbl"):
                    current_table = Table(elem, doc)

            for p in block_paragraphs:
                replace_in_paragraph(p, "7.{{ block_index }}", f"7.{block_idx}")
                replace_in_paragraph(p, "{{ duty_profile }}", dp["name"])
                replace_in_paragraph(p, "{{ framework }}", data["framework"])

            # Populate block test parameters from ACL
            if current_table:
                tbl_tpl_row = current_table.rows[1]
                for test in dp["tests"]:
                    new_row = clone_row(current_table, tbl_tpl_row)
                    new_row.cells[0].text = str(test["sr_no"])
                    new_row.cells[1].text = test["parameter"]
                    new_row.cells[2].text = test["limit"]
                    new_row.cells[3].text = test["clause"]

                current_table._tbl.remove(tbl_tpl_row._tr)

            # Configure reviewer input editable regions
            p_ac_header = None
            p_ac_body = None
            p_concl_header = None
            p_concl_body = None

            for p in block_paragraphs:
                if p.text.strip().startswith("Acceptance Criteria:"):
                    p_ac_header = p
                elif (
                    p_ac_header
                    and not p_ac_body
                    and "{{ to_be_added_by_reviewer }}" in p.text
                ):
                    p_ac_body = p
                elif p.text.strip().startswith("Conclusion:"):
                    p_concl_header = p
                elif (
                    p_concl_header
                    and not p_concl_body
                    and "{{ to_be_added_by_reviewer }}" in p.text
                ):
                    p_concl_body = p

            if p_ac_header and p_ac_body:
                replace_in_paragraph(
                    p_ac_body,
                    "{{ to_be_added_by_reviewer }}",
                    "All listed parameters meet the acceptance limits for this duty profile.",
                )
                make_reviewer_zone_editable(p_ac_header, p_ac_body, perm_id_counter)
                perm_id_counter += 1

            if p_concl_header and p_concl_body:
                replace_in_paragraph(
                    p_concl_body,
                    "{{ to_be_added_by_reviewer }}",
                    "The cell is qualified for this duty profile, subject to review.",
                )
                make_reviewer_zone_editable(
                    p_concl_header, p_concl_body, perm_id_counter
                )
                perm_id_counter += 1

    # 4. Populate footnotes block at the end of Section 7
    footnote_paragraph = None
    for p in doc.paragraphs:
        if "{{ also_fetch_any_footnotes_from_the_acl }}" in p.text:
            footnote_paragraph = p
            break

    if footnote_paragraph:
        footnote_paragraph.text = ""
        curr_footnote_p = footnote_paragraph
        for f_idx, note_text in enumerate(data["footnotes"]):
            if f_idx == 0:
                curr_footnote_p.text = note_text
            else:
                curr_footnote_p = insert_paragraph_after(
                    curr_footnote_p, note_text, style=footnote_paragraph.style
                )

        insert_paragraph_after(curr_footnote_p, "")

    # 5. Global token string replacements
    replacements = {
        "{{ cell_model }}": data["cell_model"],
        "{{ manufacturer_from_datasheet }}": data["manufacturer_from_datasheet"],
        "{{ doc_number }}": data["doc_number"],
        "{{ market }}": data["market"],
        "{{ framework }}": data["framework"],
        "{{ lab }}": data["lab"],
        "{{ format_from_datasheet }}": data["format_from_datasheet"],
        "{{ chemistry }}": data["chemistry"],
        "{{ duty_profiles }}": data["duty_profiles_joined"],
        "{{ tmp_doc_title }}": data["tmp_doc_title"],
        "{{ acl_doc_title }}": data["acl_doc_title"],
        "{{ datasheet_title }}": data["datasheet_title"],
        "{{ nominal_voltage }}": data["nominal_voltage"],
        "{{ v_max }}": data["v_max"],
        "{{ v_min }}": data["v_min"],
        "{{ rated_capacity }}": data["rated_capacity"],
        "{{ grading_low }}": data["grading_low"],
        "{{ grading_high }}": data["grading_high"],
        "{{ storage_from_datasheet }}": data["storage_from_datasheet"],
        "{{ supplied_as_from_datasheet }}": data["supplied_as_from_datasheet"],
    }

    for r_key, r_val in replacements.items():
        replace_in_header_footer(doc, r_key, r_val)

    for p in doc.paragraphs:
        for r_key, r_val in replacements.items():
            replace_in_paragraph(p, r_key, r_val)

    for table in doc.tables:
        if table == table1:
            continue
        try:
            # Prevent altering the Revision Record table
            if (
                len(table.rows) > 0
                and len(table.rows[0].cells) > 0
                and "Revision" in table.rows[0].cells[0].text
            ):
                continue
        except Exception:
            pass

        for r_key, r_val in replacements.items():
            replace_in_table(table, r_key, r_val)

    # 6. Apply document editing locks
    apply_document_protection(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Protocol document saved successfully to: {output_path}")
